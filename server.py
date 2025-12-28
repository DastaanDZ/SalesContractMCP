from fastmcp import FastMCP, Context
from pydantic import BaseModel, Field
from typing import Annotated
import os
import io
import json
import re
from dotenv import load_dotenv
from supabase import create_client, Client
from docx import Document

load_dotenv()

# --------------------------------
# 1. SETUP
# --------------------------------
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
BUCKET_NAME = "od-files"

try:
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
except Exception as e:
    print(f"⚠️  Supabase Init Error: {e}")

mcp = FastMCP(name="od-smart-drafter", version="6.0.0")

# --------------------------------
# 2. DATA MODELS
# --------------------------------
class LineItemInput(BaseModel):
    item_name: str = Field(..., description="Name of the service/product")
    description: str = Field(..., description="Short details")
    price: str = Field(..., description="Price (e.g. '$500.00')")

# --------------------------------
# 3. SMART VERSIONING HELPERS
# --------------------------------

DATA_DIR = os.path.join(os.path.dirname(__file__), "data")
CLAUSES_FILE = os.path.join(DATA_DIR, "clauses.json")

def load_clauses():
    if not os.path.exists(CLAUSES_FILE): return {}
    with open(CLAUSES_FILE, "r") as f: return json.load(f)

def get_latest_file_content(quote_number: str):
    """
    Scans Supabase for all versions of the quote.
    Returns: (latest_filename, file_stream_bytes)
    """
    try:
        # List all files
        files = supabase.storage.from_(BUCKET_NAME).list()
        all_names = [f['name'] for f in files]
        
        # Filter for files belonging to this quote
        # Matches: "100.docx", "100_v1.docx", "100_v10.docx"
        pattern = re.compile(rf"^{re.escape(quote_number)}(_v(\d+))?\.docx$")
        
        candidates = []
        for name in all_names:
            match = pattern.match(name)
            if match:
                version_num = int(match.group(2)) if match.group(2) else 0
                candidates.append((version_num, name))
        
        if not candidates:
            return None, None

        # Sort by version number descending
        candidates.sort(key=lambda x: x[0], reverse=True)
        latest_filename = candidates[0][1]
        
        # Download
        print(f"📥 Fetching latest version: {latest_filename}")
        res = supabase.storage.from_(BUCKET_NAME).download(latest_filename)
        return latest_filename, io.BytesIO(res)

    except Exception as e:
        print(f"Error fetching latest: {e}")
        return None, None

def upload_new_version(quote_number: str, stream: io.BytesIO) -> str:
    """
    Calculates next version number and uploads.
    """
    try:
        # Get current list to determine next version
        files = supabase.storage.from_(BUCKET_NAME).list()
        all_names = [f['name'] for f in files]
        
        # Calculate next version
        pattern = re.compile(rf"^{re.escape(quote_number)}(_v(\d+))?\.docx$")
        max_v = 0
        found_base = False
        
        for name in all_names:
            match = pattern.match(name)
            if match:
                found_base = True
                v = int(match.group(2)) if match.group(2) else 0
                if v > max_v: max_v = v
        
        # If we never found even the original, force v1 (though unlikely if download succeeded)
        next_v = max_v + 1
        new_filename = f"{quote_number}_v{next_v}.docx"
        
        # Upload
        stream.seek(0)
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        supabase.storage.from_(BUCKET_NAME).upload(
            path=new_filename, 
            file=stream.read(), 
            file_options={"content-type": mime, "upsert": "false"}
        )
        return supabase.storage.from_(BUCKET_NAME).get_public_url(new_filename)
        
    except Exception as e:
        raise Exception(f"Upload failed: {e}")

# --------------------------------
# 4. DUPLICATE CHECKERS
# --------------------------------

def clause_exists(doc: Document, clause_title: str) -> bool:
    """Checks if a paragraph with the clause title already exists."""
    # Check paragraphs (Headings are technically paragraphs with style)
    for p in doc.paragraphs:
        if clause_title.lower() in p.text.lower():
            return True
    return False

def row_exists(doc: Document, item_name: str, price: str) -> bool:
    """Checks if a row with matching Item Name AND Price exists."""
    if not doc.tables: return False
    
    table = doc.tables[0] # Assuming first table
    for row in table.rows:
        # Assuming Col 0 is Name, Col 2 is Price (based on previous schema)
        # We use simple string containment for safety
        cells = [c.text.lower() for c in row.cells]
        if item_name.lower() in cells[0] and price.lower() in cells[2]:
            return True
    return False

# --------------------------------
# 5. TOOLS
# --------------------------------

@mcp.tool(
    name="draft_docx_od",
    description="Append a legal clause to the latest version of the OD.",
)
async def draft_docx_od(
    ctx: Context, 
    quote_number: Annotated[str, "The ID of the quote"], 
    clause_name: Annotated[str, "Exact name of the clause"]
):
    # 1. Validate Clause Name
    clause_db = load_clauses()
    found_key = next((k for k in clause_db.keys() if k.lower() in clause_name.lower()), None)
    
    if not found_key:
        available = ", ".join(clause_db.keys())
        return {"content": [{"type": "text", "text": f"❌ Clause '{clause_name}' not found. Options: {available}"}], "isError": True}

    # 2. Get LATEST Version (Auto-resolution)
    latest_name, stream = get_latest_file_content(quote_number)
    if not stream:
        return {"content": [{"type": "text", "text": f"❌ No files found for Quote {quote_number}"}], "isError": True}

    await ctx.info(f"Editing latest file: {latest_name}")

    # 3. Edit & Check Duplicates
    try:
        doc = Document(stream)
        
        # --- IDEMPOTENCY CHECK ---
        if clause_exists(doc, found_key):
             return {
                "content": [{"type": "text", "text": f"⚠️ Clause '{found_key}' already exists in {latest_name}. No changes made."}]
            }

        # Apply Change
        doc.add_heading(found_key, level=2)
        doc.add_paragraph(clause_db[found_key])
        
        # Save & Upload
        out = io.BytesIO()
        doc.save(out)
        public_url = upload_new_version(quote_number, out)
        
        return {
            "content": [
                {"type": "text", "text": f"✅ Added '{found_key}'. New version created."},
                {"type": "text", "text": f"📄 View: {public_url}"}
            ]
        }
    except Exception as e:
        return {"content": [{"type": "text", "text": f"❌ Error: {str(e)}"}], "isError": True}


# @mcp.tool(
#     name="add_line_item",
#     description="Add a row to the pricing table in the latest OD version.",
# )
# async def add_line_item(
#     ctx: Context, 
#     quote_number: Annotated[str, "The ID of the quote"],
#     item_name: Annotated[str, "Name of service"] = None, 
#     description: Annotated[str, "Description"] = None, 
#     price: Annotated[str, "Price"] = None
# ):
#     # 1. Elicitation
#     if not all([item_name, description, price]):
#         await ctx.info(f"Triggering elicitation for Quote {quote_number}")
#         result = await ctx.elicit(
#             message=f"I need details to add a row to Quote {quote_number}.",
#             response_type=LineItemInput
#         )
#         if result.action != "accept": return "Cancelled."
#         item_name, description, price = result.data.item_name, result.data.description, result.data.price

#     # 2. Get LATEST Version
#     latest_name, stream = get_latest_file_content(quote_number)
#     if not stream:
#         return {"content": [{"type": "text", "text": f"❌ Quote {quote_number} not found."}], "isError": True}

#     await ctx.info(f"Editing latest file: {latest_name}")

#     # 3. Edit & Check Duplicates
#     try:
#         doc = Document(stream)
#         if not doc.tables: return "❌ No tables found."
        
#         # --- IDEMPOTENCY CHECK ---
#         if row_exists(doc, item_name, price):
#              return {
#                 "content": [{"type": "text", "text": f"⚠️ Row '{item_name}' ({price}) already exists in {latest_name}. No changes made."}]
#             }

#         # Add Row
#         table = doc.tables[0]
#         row = table.add_row()
#         row.cells[0].text = item_name
#         row.cells[1].text = description
#         row.cells[2].text = price

#         # Save & Upload
#         out = io.BytesIO()
#         doc.save(out)
#         public_url = upload_new_version(quote_number, out)

#         return {
#             "content": [
#                 {"type": "text", "text": f"✅ Added row '{item_name}'. New version created."},
#                 {"type": "text", "text": f"📄 View: {public_url}"}
#             ]
#         }
#     except Exception as e:
#         return {"content": [{"type": "text", "text": f"❌ Error: {str(e)}"}], "isError": True}


@mcp.tool(
    name="add_line_item",
    description="Add a row to the pricing table in the latest OD version.",
)
async def add_line_item(
    ctx: Context, 
    quote_number: Annotated[str, "The ID of the quote"],
    item_name: Annotated[str, "Name of service"] = None, 
    description: Annotated[str, "Description"] = None, 
    price: Annotated[str, "Price"] = None
):
    # --- 1. ROBUST PARAMETER CHECK (Replaces Elicitation) ---
    # Since Cline doesn't support 'elicit', we just return a helpful message
    # if not all([item_name, description, price]):
    #     missing = []
    #     if not item_name: missing.append("item_name")
    #     if not description: missing.append("description")
    #     if not price: missing.append("price")
        
    #     return f"⚠️ Missing details: {', '.join(missing)}. Please try again providing Item Name, Description, and Price."

    # 2. Get File
    latest_name, stream = get_latest_file_content(quote_number)
    if not stream:
        return f"❌ Quote {quote_number} not found."

    await ctx.info(f"Adding row '{item_name}' to {latest_name}")

    # 3. Edit
    try:
        doc = Document(stream)
        if not doc.tables: return "❌ No tables found in document."
        
        if row_exists(doc, item_name, price):
             return f"⚠️ Row '{item_name}' already exists in {latest_name}. No changes made."

        table = doc.tables[0]
        row = table.add_row()
        row.cells[0].text = item_name
        row.cells[1].text = description
        row.cells[2].text = price

        out = io.BytesIO()
        doc.save(out)
        public_url = upload_new_version(quote_number, out)

        return f"✅ Added row '{item_name}'. View: {public_url}"
    except Exception as e:
        return f"❌ Error: {str(e)}"
    
    
if __name__ == "__main__":
    mcp.run(transport="http", port=8000)